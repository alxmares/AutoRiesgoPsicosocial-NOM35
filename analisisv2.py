import numpy as np
import matplotlib
import matplotlib.pyplot as plt
from tkinter import filedialog
import threading
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import seaborn as sns
import threading
import webbrowser
from dash import Dash, dcc, html
import plotly.graph_objs as go
import plotly.colors as pcolors  # Importar el módulo de colores de Plotly
import pdfkit
import tempfile
import os

# Asegurarse que no utilice backend de CTkinter
matplotlib.use('Agg')

class Analisis():
    def __init__(self, callback_function1, callback_function2):
        self.callback_function1 = callback_function1
        self.callback_function2 = callback_function2
        
        self.categorias = {
            "index": {1:[1,2,3,4,5], 2:[6,7,8,9,10,11,12,13,14,15,16,23,24,25,26,27,28,29,30,35,36,65,66,67,68],
                      3:[17,18,19,20,21,22], 4:[31, 32, 33, 34, 37, 38, 39, 40, 41, 42, 43, 44, 45, 46, 57, 58, 59, 60, 61, 62, 63, 64, 69, 70, 71, 72],
                      5:[47, 48, 49, 50, 51, 52, 53, 54, 55, 56]},
            "range":{1:[5,9,11,14],2:[15,30,45,60],3:[5,7,10,13],4:[14,29,42,58],5:[10,14,18,23]}
        }
        
        self.dominios = {
            "index": {1:[1,3,2,4,5],2:[6,12,7,8,9,10,11,65,66,67,68,13,14,15,16],
                      3:[23, 24, 29,30,35,36],4:[17,18],5:[19,20,21,22],6:[31, 32, 33, 34,37, 38, 39, 40, 41],
                      7:[42, 43, 44, 45, 46, 69, 70, 71, 72],8:[57, 58, 59, 60, 61, 62, 63, 64],
                      9:[47, 48, 49, 50, 51, 52],10:[55, 56, 53, 54]},
            "range":{1:[5,9,11,14],2:[15,21,27,37],3:[11,16,21,25],
                     4:[1,2,4,6],5:[4,6,8,10],6:[9,12,16,20],7:[10,13,17,21],
                     8:[7,10,13,14,16],9:[6,10,14,18],10:[4,6,8,10]}
        }
        
        self.rango_final = [(0,50),(50,75),(75,99),(99,140),(140,300)]
        
        self.int_calificaciones = [1,2,3,4,5]
        self.calificaciones = ["Nulo", "Bajo", "Medio", "Alto", "Muy Alto"]
        self.calif_der = [1,4,23,24,25,26,27,28,30,31,32,33,34,35,36,37,38,39,40,41,42,43,44,45,46,47,48,49,50,51,52,53,55,56,57]
        self.calif_izq = [2,3,5,6,7,8,9,10,11,12,13,14,15,16,17,18,19,20,21,22,29,54,58,59,60,61,62,63,64,65,66,67,68,69,70,71,72]
        
        self.colores_por_calificacion = ['#1f77b4', '#2ca02c', '#ffd700', '#ff7f0e', '#d62728']
        
        self.necesidad_accion = {
            "Muy Alto": "Es preciso realizar el análisis de cada categoría y dominio para establecer las acciones de intervención apropiadas, mediante un Programa de intervención intensivo que deberá incluir campañas de sensibilización, aplicar la política de prevención de riesgos psicosociales e implementar programas para la prevención de los factores de riesgo psicosocial, la promoción de un entorno organizacional favorable y la prevención de la violencia laboral, así como garantizar su aplicación y difusión.",
            "Alto":"Es preciso realizar un análisis de cada categoría y dominio, de manera que se puedan determinar las acciones de intervención concretas a través de un Programa de intervención, que podrá incluir una evaluación específica y deberá incluir una campaña de sensibilización, implementar la política de prevención de riesgos psicosociales y programas para la prevención de los factores de riesgo psicosocial, la promoción de un entorno organizacional favorable y la prevención de la violencia laboral, así como garantizar su aplicación y difusión.",
            "Medio":"Se requiere reforzar la política de prevención de riesgos psicosociales y programas para la prevención de los factores de riesgo psicosocial, la promoción de un entorno organizacional favorable y la promoción de un entorno laboral saludable, así como garantizar su aplicación y difusión, mediante un Programa de intervención.",
            "Bajo":"Es necesario una mayor difusión de la política de prevención de riesgos psicosociales y programas para: la prevención de los factores de riesgo psicosocial, la promoción de un entorno organizacional favorable y la prevención de la violencia laboral.",
            "Nulo":"El riesgo resulta poco significativo por lo que no se requiere medidas adicionales.",  
        }
        self.cat_nombres = ['Ambiente de trabajo', 'Factores propios de actividad', 
                      'Organización tiempo de trabajo', 'Liderazgo relaciones en trabajo',
                      'Entorno organizacional']
        self.dom_nombres = ['Condiciones ambiente de trabajo', 'Cargo de trabajo', 'Falta de control trabajo', 
                      'Jornada de trabajo','Interferencia relación trabajo-familia','Liderazgo',
                      'Relaciones en trabajo','Violencia','Reconocimiento desempeño',
                      'Insuficiente sentido pertenecencia inestabilidad']
        
    def analizar(self, nombre, edad, empresa, puesto, respuestas):
        self.nombre = nombre
        self.edad = edad
        self.empresa = empresa
        self.puesto = puesto
        self.puntajes = respuestas
        
        if len(self.calificaciones)>5:
            self.calificaciones.pop(0)
            
        self.resultado_cat = {"punt":{1:0,2:0,3:0,4:0,5:0},
                         "cal":{1:0,2:0,3:0,4:0,5:0}}
        
        self.resultado_dom = {"punt":{1:0,2:0,3:0,4:0,5:0,6:0,7:0,8:0,9:0,10:0},
                         "cal":{1:"",2:"",3:"",4:"",5:"",6:"",7:"",8:"",9:"",10:""}}
        self.punt_final = 0
        # categoría
        for idx,respuesta in enumerate(respuestas):
            idx=idx+1 # índices comienzan en 1
            if respuesta == "":continue # De las 65 a 72 no importa si responde o no 
            
            # Buscar categoria
            for i in range(1,6):
                if idx in self.categorias["index"][i]:
                    categoria = i
            
            # Buscar dominio
            for i in range(1,11):
                if idx in self.dominios["index"][i]:
                    dominio = i 
            
            # Obtener puntuación
            if idx in self.calif_der:
                punt = respuesta
            else:
                punt = 4-respuesta
            
            # Sumar calificaciones
            self.punt_final += punt
            self.resultado_cat["punt"][categoria] += punt
            self.resultado_dom["punt"][dominio] += punt
        
        # Obtener calificación de categorias
        for i in range(1,6):
            puntuacion = self.resultado_cat["punt"][i]
            rangos = self.categorias["range"][i]
            if puntuacion < rangos[0]:
                calificacion = self.int_calificaciones[0]
            elif puntuacion < rangos[1]:
                calificacion = self.int_calificaciones[1]
            elif puntuacion < rangos[2]:
                calificacion = self.int_calificaciones[2]
            elif puntuacion < rangos[3]:
                calificacion = self.int_calificaciones[3]
            else:
                calificacion = self.int_calificaciones[4]
            
            self.resultado_cat["cal"][i]=calificacion
            
        # Obtener calificación de dominio
        for i in range(1,11):
            puntuacion = self.resultado_dom["punt"][i]
            rangos = self.dominios["range"][i]
            if puntuacion < rangos[0]:
                calificacion = self.int_calificaciones[0]
            elif puntuacion < rangos[1]:
                calificacion = self.int_calificaciones[1]
            elif puntuacion < rangos[2]:
                calificacion = self.int_calificaciones[2]
            elif puntuacion < rangos[3]:
                calificacion = self.int_calificaciones[3]
            else:
                calificacion = self.int_calificaciones[4]
            
            self.resultado_dom["cal"][i]=calificacion
            
        # Obtener calificación final
        for i,(rango,calf) in enumerate(zip(self.rango_final, self.calificaciones)):
            if rango[0] <= self.punt_final < rango[1]:
                self.cal_final = calf
                self.color_final = self.colores_por_calificacion[i]
                #print(self.cal_final)
                #print(self.color_final) 
                break
        
        self.accion = self.necesidad_accion[self.cal_final]
        #print(f'Puntaje_final: {self.punt_final}, Calificación: {self.cal_final}')
        #print(self.resultado_cat)
        #print(self.resultado_dom)
        
        return self.punt_final,self.cal_final
    
    def graficar(self):
        # Crear un hilo para ejecutar el servidor Dash
        dash_thread = threading.Thread(target=self.run_dash)
        dash_thread.daemon = True  # Hacer que el hilo termine cuando se cierra la app
        dash_thread.start()

    def run_dash(self):
        # Preparar datos para los gráficos
        calificaciones_categorias = list(self.resultado_cat["cal"].values())
        nombres_categorias = self.cat_nombres
        calificaciones_dominios = list(self.resultado_dom["cal"].values())
        nombres_dominios = self.dom_nombres

        print(calificaciones_categorias, calificaciones_dominios)
        # Crear aplicación Dash
        app = Dash(__name__)

        app.layout = html.Div([
            html.H1(f"{self.nombre} : {self.cal_final}"),

            # Gráfico de Barras - Calificaciones por Categoría
            html.Div([
                dcc.Graph(
                    id='grafico_barras_categorias',
                    figure=self.crear_grafico_barras(calificaciones_categorias, nombres_categorias, 'Calificaciones por Categoría')
                )
            ], style={'width': '50%', 'display': 'inline-block'}),

            # Gráfico de Pastel - Distribución de Calificaciones por Categoría
            html.Div([
                dcc.Graph(
                    id='grafico_pastel_categorias',
                    figure=self.crear_grafico_pastel(calificaciones_categorias, nombres_categorias, 'Distribución por Categoría')
                )
            ], style={'width': '50%', 'display': 'inline-block'}),

            # Gráfico de Barras - Calificaciones por Dominio
            html.Div([
                dcc.Graph(
                    id='grafico_barras_dominios',
                    figure=self.crear_grafico_barras(calificaciones_dominios, nombres_dominios, 'Calificaciones por Dominio')
                )
            ], style={'width': '50%', 'display': 'inline-block'}),

            # Gráfico de Pastel - Distribución de Calificaciones por Dominio
            html.Div([
                dcc.Graph(
                    id='grafico_pastel_dominios',
                    figure=self.crear_grafico_pastel(calificaciones_dominios, nombres_dominios, 'Distribución por Dominio')
                )
            ], style={'width': '50%', 'display': 'inline-block'})
        ])

        # Abrir el navegador automáticamente
        webbrowser.open_new("http://127.0.0.1:8050")

        # Iniciar el servidor Dash
        app.run_server(debug=False, use_reloader=False)

    def crear_grafico_barras(self, calificaciones, nombres, titulo):
    # Crear un gráfico de barras con una escala de colores basada en los valores de calificación
        fig = go.Figure(
            data=[
                go.Bar(
                    x=nombres,
                    y=list(calificaciones),  # Convertir dict_values a una lista
                    marker=dict(
                        color=list(calificaciones),  # Utilizar los valores como base para el color
                        colorscale='Viridis',  # Aplicar la paleta de colores 'Viridis'
                        cmin=1,  # Mínimo valor para la escala de color
                        cmax=5,  # Máximo valor para la escala de color
                    )
                )
            ],
            layout=go.Layout(
                title=titulo,
                xaxis=dict(title='Categorías/Dominios'),
                yaxis=dict(title='Calificación'),
                template='plotly_white'  # Estilo minimalista con fondo blanco
            )
        )
        
        return fig

    def crear_grafico_pastel(self, calificaciones, nombres, titulo):
        fig = go.Figure(
            data=[
                go.Pie(
                    labels=nombres,
                    values=calificaciones,
                    hole=.3,
                    marker=dict(colors=pcolors.sequential.Viridis),  # Corregido para usar pcolors
                )
            ],
            layout=go.Layout(
                title=titulo,
                template='plotly_white'  # Estilo minimalista con fondo blanco
            )
        )

        return fig
    
    def crear_grafico_barras_matplotlib(self, calificaciones, nombres, titulo):
        """
        Crea un gráfico de barras usando Matplotlib y guarda la gráfica como PNG.
        """
        # Crear una figura
        fig, ax = plt.subplots(figsize=(10, 6))

        # Crear gráfico de barras
        ax.bar(nombres, calificaciones, color=plt.cm.viridis(np.linspace(0, 1, len(calificaciones))))
        
        # Configurar título y etiquetas
        ax.set_title(titulo)
        ax.set_xlabel('Categorías/Dominios')
        ax.set_ylabel('Calificación')
        ax.grid(True)
        # Rotar los ticks del eje x a 45 grados
        ax.set_xticklabels(nombres, rotation=45, ha='right')


        # Guardar la gráfica como PNG usando la figura creada
        self.guardar_grafica_matplotlib(fig, 'grafico_barras_matplotlib.png')

        # Retornar la figura creada
        return fig

    def guardar_grafica_matplotlib(self, fig, filename, formato='png'):
        """
        Guarda una gráfica de Matplotlib en el formato especificado.
        """
        try:
            fig.savefig(filename, format=formato)
            print(f'Gráfica guardada como {filename}')
        except Exception as e:
            print(f'Error al guardar gráfica: {e}')
        finally:
            plt.close(fig)  # Cerrar la figura para liberar memoria

    def hex_to_rgb(self, hex_color):
        hex_color = hex_color.lstrip('#')
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    
    def save_doc(self):
        """
        Guarda el archivo .docx. En caso de no hacerlo retorna None
        """
        # Guardar documento
        dirname = self.nombre.replace(" ", "_")
        filename = "resultados_"+dirname+".docx"
        
        opciones = {
        'defaultextension': '.docx',
        'filetypes': [('Documentos de Word', '.docx'), ('Todos los archivos', '.*')],
        'title': 'Guardar como...',
        "initialfile": filename
        }
    
        filename = filedialog.asksaveasfilename(**opciones)
        if filename:
            #print(f"Archivo guardado en {filename}")
            return filename,dirname
        
        return None, None
    
    def search_logo(self):
        filepath = filedialog.askopenfilename(
          title="Selecciona una imagen para el encabezado",
        filetypes=[("Archivos de imagen", "*.png;*.jpg;*.jpeg;*.gif")]
        )

        # Verificar si se seleccionó un archivo
        if not filepath:
            #print("No se seleccionó ningún archivo.")
            return
        
        return filepath
    
    def hilo_word(self, logo_fp):
        """
        Se crea el archivo de Word en un hilo.
        """

        filename,dirname = self.save_doc() # Obtener locación del archivo, no lo guarda
        
        # Significa que canceló la acción
        if not filename or not dirname:
            return 
        
        info = {
            "Cal": self.cal_final,
            "Nombre": self.nombre,
            "Edad": self.edad,
            "Empresa": self.empresa,
            "Puesto": self.puesto,
            "Puntaje Final": self.punt_final,
            "Calificación final": self.cal_final,
            "Necesidad de acción": self.accion
            }
        
        doc = Document()
        # Acceder a la sección actual del documento para ajustar los márgenes
        section = doc.sections[0]
        section.top_margin = Inches(1)  # Margen superior de 1 pulgada
        section.bottom_margin = Inches(1)  # Margen inferior de 1 pulgada
        section.left_margin = Inches(1)  # Margen izquierdo de 1 pulgada
        section.right_margin = Inches(1)  # Margen derecho de 1 pulgada
        
        # Establecer la fuente predeterminada del documento
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Añadir logotipo en caso de que se haya declarado
        if logo_fp:
            # Acceder al encabezado de la sección por defecto del documento
            section = doc.sections[0]
            header = section.header

            # Añadir un párrafo al encabezado y alinear el texto a la derecha
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.alignment = 2  # 2 es el valor para la alineación a la derecha

            # Añadir la imagen al párrafo del encabezado
            run = paragraph.add_run()
            run.add_picture(logo_fp, width=Inches(1))  # Ajusta el tamaño según sea necesario

        # = = =  H E A D I N G  = = =
        heading = doc.add_heading(level=0)
        h = heading.add_run("Reporte de Evaluación: ")
        h.font.size = Pt(22)
        h.font.name = "Calibri"
        h = heading.add_run(self.cal_final)
        r,g,b = self.hex_to_rgb(self.color_final)
        h.font.color.rgb = RGBColor(r,g,b)
        h.font.size = Pt(22)
        h.font.name = "Calibri"
        # ===========================

        # Añadir un párrafo al documento
        p = doc.add_paragraph()
        # Establecer el espaciado de línea
        p.paragraph_format.line_spacing = 1.5  # Espaciado de línea 1.5

        # Establecer el espacio después del párrafo (espacio entre líneas)
        p.paragraph_format.space_after = Pt(6)
        # Indicar la clave específica donde quieres agregar un salto de línea
        salto_de_linea = ["Nombre"]

        # Iterar sobre los elementos del diccionario y agregarlos al párrafo
        for key, value in info.items():
            if key == "Cal":
                continue  # Ignorar esta clave
            
            p.add_run(f"{key}: ").bold=True
            p.add_run(f"{value}")
            
            if key == "Puntaje Final":
                break
            # Comprobar si es la clave para insertar un salto de línea
            if key in salto_de_linea:
                # Agregar una tabulación, se asume que se quiere doble tabulación
                p.add_run('\t\t')
            else:
                p.add_run().add_break()  # Insertar salto de línea



        p = doc.add_paragraph()
        p.alignment= WD_ALIGN_PARAGRAPH.JUSTIFY
        p.add_run("Necesidad de acción: ").bold=True
        p.add_run(info["Necesidad de acción"])
        
        # Insertar las cuatro gráficas en una tabla
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
         # Insertar la gráfica de barras en una tabla
        table = doc.add_table(rows=2, cols=1)  # Cambiado a una sola celda
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        


        fig_barras_categoria = self.crear_grafico_barras_matplotlib(
            self.resultado_cat["cal"].values(), self.cat_nombres, 'Calificaciones por Categoría'
        )

        img_stream = io.BytesIO()
        fig_barras_categoria.savefig(img_stream, format='png')
        img_stream.seek(0)
        cell = table.cell(0, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_stream, width=Inches(6.5))  # Ajustar al ancho del documento
        paragraph.add_run("\nGráfico de Barras por Categoría").bold = True

        # Generar gráfica de barras por Dominio
        fig_barras_dominio = self.crear_grafico_barras_matplotlib(
            self.resultado_dom["cal"].values(), self.dom_nombres, 'Calificaciones por Dominio'
        )

        img_stream = io.BytesIO()
        fig_barras_dominio.savefig(img_stream, format='png')
        img_stream.seek(0)
        cell = table.cell(1, 0)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_stream, width=Inches(6.5))  # Ajustar al ancho del documento
        paragraph.add_run("\nGráfico de Barras por Dominio").bold = True


        # Añadir espacio para la firma
        p = doc.add_paragraph()
        run = p.add_run("\n\n\nFirma de conformidad:\n\n\n").bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("_" * 30).bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Nombre, fecha y firma").bold = True

        # Guardar el documento como Word
        try:
            doc.save(filename)
            print(f"Documento guardado como {filename}")
            self.callback_function2()
        except Exception as e:
            print(f"Error al guardar documento: {e}")
            self.callback_function1()
    
    def generar_documento(self, op):
        # Buscar logotipo
        if op:
            logo_fp = self.search_logo()
        else: 
            logo_fp = False
            
        word_thread = threading.Thread(target=self.hilo_word, args=(logo_fp,))
        word_thread.daemon = True
        word_thread.start()